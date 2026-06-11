"""
Bid Compliance Report Generator — CEO-Ready Format
Matches the professional style of BESS_Tender_Contents_Checklist reference document.

Usage:
    python generate_report.py
    python generate_report.py --timestamp 20260611_153155
    python generate_report.py --output reports/my_report.docx
"""

import json
import argparse
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config.settings import (
    METADATA_OUTPUT_DIR, COMPLIANCE_OUTPUT_DIR,
    OEM_OUTPUT_DIR, ENVELOPE_OUTPUT_DIR, OUTPUTS_DIR
)

# ─── Colour palette (extracted from reference document) ──────────────────────
C_NAVY       = "1F3864"   # dark navy  — main title, general section headers
C_BLUE_HDR   = "1F497D"   # medium blue — Envelope-01 headers, table column headers
C_BLUE_MED   = "2E75B6"   # bright blue — subtitle
C_GREEN_HDR  = "375623"   # dark green  — Envelope-02 headers
C_RED        = "C00000"   # deep red    — deadline, warnings, Critical risk
C_ORANGE     = "E36C09"   # orange      — High risk
C_AMBER      = "7F6000"   # dark amber  — Medium risk
C_GREY_TXT   = "444444"   # dark grey   — project name text
C_GREY_NOTE  = "666666"   # light grey  — footer notes
C_WHITE      = "FFFFFF"
C_ROW_ODD    = "FFFFFF"
C_ROW_EVEN   = "F2F2F2"
C_SEC_ROW    = "D6E4F7"   # light blue  — section group rows inside tables
C_MAND_FILL  = "FCE4D6"   # salmon      — MANDATORY status cell
C_COND_FILL  = "FBE4D5"   # light peach — CONDITIONAL status cell
C_OPT_FILL   = "E2EFDA"   # light green — OPTIONAL status cell
C_ENV1_HDR   = "D6E4F7"   # same as section rows — Env-01 cross-ref
C_ENV2_HDR   = "FBE4D5"   # Env-02 cross-ref


# ─── XML helpers ─────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color="CCCCCC", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(existing)
    mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _cell_write(cell, text: str, bold=False, italic=False,
                font_size=9, color=C_NAVY, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Clear cell and write a single-run paragraph."""
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)


def _cell_write_multi(cell, lines: list[tuple], font_size=8.5):
    """Write multiple runs/lines into a cell (list of (text, bold, color))."""
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    for text, bold, color in lines:
        run = para.add_run(str(text) if text else "")
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(color)


def _para(doc, text, bold=False, italic=False, font_size=10,
          color=C_NAVY, align=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def _section_heading(doc, text, color=C_NAVY, font_size=14, space_before=16):
    return _para(doc, text, bold=True, font_size=font_size,
                 color=color, space_before=space_before, space_after=5)


def _warning(doc, text):
    return _para(doc, f"⚠  {text}", bold=True, font_size=8.5,
                 color=C_RED, space_before=3, space_after=3)


def _risk_color(risk: str) -> str:
    r = (risk or "").lower()
    if "critical" in r:
        return C_RED
    if "high" in r:
        return C_ORANGE
    if "medium" in r:
        return C_AMBER
    return C_GREY_TXT


def _infer_risk(doc_name: str, is_mandatory: bool, category: str = "") -> str:
    name_lower = (doc_name or "").lower()
    cat = (category or "").upper()
    keywords_critical = [
        "security", "bid security", "submission letter", "letter of bid",
        "authorization", "authorisation", "performance", "compliance",
        "price schedule", "financial offer", "technical offer",
        "completion certif", "specific experience", "manufacturer"
    ]
    if any(k in name_lower for k in keywords_critical):
        return "Critical"
    if "MANDATORY" in cat or is_mandatory:
        return "High"
    if "CONDITIONAL" in cat:
        return "Medium"
    return "Low"


def _status_fill(status: str) -> str:
    s = (status or "").upper()
    if "MANDATORY" in s:
        return C_MAND_FILL
    if "CONDITIONAL" in s:
        return C_COND_FILL
    return C_OPT_FILL


def _status_color(status: str) -> str:
    s = (status or "").upper()
    if "MANDATORY" in s:
        return C_RED
    if "CONDITIONAL" in s:
        return C_ORANGE
    return "375623"


# ─── Data loaders ─────────────────────────────────────────────────────────────

def _load(directory: Path, timestamp: str | None) -> dict:
    if timestamp:
        matches = list(directory.glob(f"{timestamp}_*.json"))
        if matches:
            return json.loads(matches[0].read_text(encoding="utf-8"))
    files = sorted(directory.glob("*.json"))
    return json.loads(files[-1].read_text(encoding="utf-8")) if files else {}


# ─── Title / Header ──────────────────────────────────────────────────────────

def _add_title_block(doc: Document, meta: dict):
    # Main title
    _para(doc, "BID COMPLIANCE ANALYSIS REPORT",
          bold=True, font_size=22, color=C_NAVY,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

    # Subtitle — method
    method = meta.get("procurement_method") or "One Stage Two Envelope Tendering Method"
    _para(doc, method,
          bold=True, font_size=13, color=C_BLUE_MED,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=3)

    # Project name
    project = meta.get("project_name") or ""
    if project:
        _para(doc, project, font_size=10, color=C_GREY_TXT,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)

    # "Complete Contents Checklist for:"
    _para(doc, "Complete Bid Compliance Checklist for:",
          bold=True, font_size=10, color=C_NAVY,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=3)

    # Deadline in red
    deadline = meta.get("bid_submission_deadline") or ""
    if deadline:
        _para(doc, f"Submission Deadline:  {deadline}",
              bold=True, font_size=11, color=C_RED,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=3)

    # Package / authority / funding
    parts = []
    ref = meta.get("tender_reference") or ""
    entity = meta.get("procuring_entity") or ""
    fund = meta.get("funding_source") or ""
    if ref:
        parts.append(f"Tender Ref: {ref}")
    if entity:
        parts.append(f"Issuing Authority: {entity}")
    if fund:
        parts.append(f"Funding: {fund}")
    if parts:
        _para(doc, "  |  ".join(parts), font_size=8.5, color=C_GREY_NOTE,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=8)


# ─── Section 1: Submission Format Overview ───────────────────────────────────

def _add_submission_overview(doc: Document, meta: dict, env: dict):
    _section_heading(doc, "SUBMISSION FORMAT OVERVIEW", color=C_NAVY)

    sub_fmt = env.get("submission_format") or ""
    if sub_fmt:
        _para(doc, sub_fmt[:500], font_size=9, color="000000",
              space_before=2, space_after=4)

    # Summary 2-col box
    e1_count = len(env.get("envelope_1", []))
    e2_count = len(env.get("envelope_2", []))

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = tbl.rows[0]
    for cell, text, bg in [
        (row.cells[0], f"ENVELOPE-01  TECHNICAL OFFER  ({e1_count} Items)", C_BLUE_HDR),
        (row.cells[1], f"ENVELOPE-02  FINANCIAL OFFER  ({e2_count} Items)", "375623"),
    ]:
        _set_cell_bg(cell, bg)
        _set_cell_borders(cell, "FFFFFF", "6")
        _set_cell_margins(cell, 100, 100, 150, 150)
        _cell_write(cell, text, bold=True, font_size=10.5,
                    color=C_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Parameters table
    e1_mark = (env.get("envelope_1_marking") or "")[:200]
    e2_mark = (env.get("envelope_2_marking") or "")[:200]
    seal = (env.get("sealing_requirements") or "")[:200]

    params = [
        ("Copies Required",    "1 Original + 2 Copies (as specified per document)",
                               "1 Original + 2 Copies (as specified per document)"),
        ("Envelope Labelling", e1_mark or "ENVELOPE-01: TECHNICAL OFFER",
                               e2_mark or "ENVELOPE-02: FINANCIAL OFFER"),
        ("Sealing",            seal or "All envelopes must be properly sealed",
                               "Sealed; not to be opened until Technical evaluation approved"),
        ("Bid Validity",       meta.get("bid_validity_period") or "—",
                               "As per Technical Offer"),
        ("Bid Security",       meta.get("bid_security_amount") or "—",
                               "N/A — placed in Technical Offer only"),
        ("Contract Duration",  meta.get("contract_duration") or "—",
                               "—"),
    ]

    ptbl = doc.add_table(rows=0, cols=3)
    ptbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = ptbl.add_row()
    for cell, text in zip(hdr.cells, ["Parameter", "ENVELOPE-01: TECHNICAL OFFER", "ENVELOPE-02: FINANCIAL OFFER"]):
        _set_cell_bg(cell, C_BLUE_HDR)
        _set_cell_borders(cell, "FFFFFF", "4")
        _set_cell_margins(cell)
        _cell_write(cell, text, bold=True, font_size=8.5, color=C_WHITE,
                    align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, (param, v1, v2) in enumerate(params):
        row = ptbl.add_row()
        bg = C_ROW_EVEN if i % 2 == 0 else C_ROW_ODD
        _set_cell_bg(row.cells[0], "EEF2F7")
        _set_cell_bg(row.cells[1], bg)
        _set_cell_bg(row.cells[2], bg)
        for cell in row.cells:
            _set_cell_borders(cell)
            _set_cell_margins(cell)
        _cell_write(row.cells[0], param, bold=True, font_size=8.5, color=C_NAVY)
        _cell_write(row.cells[1], v1, font_size=8.5, color="000000")
        _cell_write(row.cells[2], v2, font_size=8.5, color="000000")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ─── Section 2: Tender Metadata Sheet ────────────────────────────────────────

def _add_metadata_section(doc: Document, meta: dict):
    doc.add_paragraph()  # spacer
    _section_heading(doc, "TENDER METADATA SHEET", color=C_NAVY)

    fields = [
        ("Project Name",               "project_name"),
        ("Tender Reference No.",        "tender_reference"),
        ("Procuring Entity / Employer", "procuring_entity"),
        ("Country",                     "country"),
        ("Project Location",            "project_location"),
        ("Procurement Method",          "procurement_method"),
        ("Procurement Type",            "procurement_type"),
        ("Funding Source",              "funding_source"),
        ("Loan / Credit Number",        "loan_credit_number"),
        ("Estimated Contract Value",    "estimated_contract_value"),
        ("Currency",                    "currency"),
        ("Bid Submission Deadline",     "bid_submission_deadline"),
        ("Bid Opening Date",            "bid_opening_date"),
        ("Bid Validity Period",         "bid_validity_period"),
        ("Bid Security Amount",         "bid_security_amount"),
        ("Performance Security",        "performance_security"),
        ("Advance Payment",             "advance_payment"),
        ("Contract Duration",           "contract_duration"),
        ("Pre-Bid Meeting",             "pre_bid_meeting_date"),
        ("Document Issuance Date",      "document_issuance_date"),
        ("Qualification Criteria",      "qualification_criteria_summary"),
        ("Lot Structure",               "lots"),
        ("Contact Details",             "contact_details"),
    ]

    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = tbl.add_row()
    for cell, text in zip(hdr.cells, ["Field", "Value"]):
        _set_cell_bg(cell, C_BLUE_HDR)
        _set_cell_borders(cell, "FFFFFF", "4")
        _set_cell_margins(cell)
        _cell_write(cell, text, bold=True, font_size=8.5, color=C_WHITE)

    for i, (label, key) in enumerate(fields):
        val = meta.get(key)
        if val is None or val == "":
            val = "—"
        if isinstance(val, (list, dict)):
            val = json.dumps(val, ensure_ascii=False)
        bg = C_ROW_EVEN if i % 2 == 0 else C_ROW_ODD
        row = tbl.add_row()
        _set_cell_bg(row.cells[0], "EEF2F7")
        _set_cell_bg(row.cells[1], bg)
        for cell in row.cells:
            _set_cell_borders(cell)
            _set_cell_margins(cell)
        _cell_write(row.cells[0], label, bold=True, font_size=8.5, color=C_NAVY)
        _cell_write(row.cells[1], str(val), font_size=8.5, color="000000")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ─── Envelope checklist table (shared) ────────────────────────────────────────

def _add_envelope_table(doc: Document, items: list, env_num: int):
    """6-column checklist: SL | Document / Item | Detail Requirements & Notes | Reference | Status | Rejection Risk"""
    if not items:
        _para(doc, "No documents specified.", font_size=9, color=C_GREY_TXT)
        return

    hdr_fill = C_BLUE_HDR if env_num == 1 else "2D6A4F"

    tbl = doc.add_table(rows=0, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Column header row
    hdr = tbl.add_row()
    for cell, text in zip(hdr.cells,
                          ["SL", "Document / Item", "Detail Requirements & Notes",
                           "Reference / Format", "Status", "Rejection\nRisk"]):
        _set_cell_bg(cell, hdr_fill)
        _set_cell_borders(cell, "FFFFFF", "4")
        _set_cell_margins(cell, 60, 60, 80, 80)
        _cell_write(cell, text, bold=True, font_size=8.5, color=C_WHITE,
                    align=WD_ALIGN_PARAGRAPH.CENTER)

    # Group items by section
    groups: dict[str, list] = {}
    for item in items:
        # Try to detect a section group from doc_id pattern or notes
        doc_id = item.get("doc_id", "")
        # Use first word of document name as group heuristic, or notes
        grp = _detect_section(item, env_num)
        groups.setdefault(grp, []).append(item)

    row_num = 1
    for section_name, section_items in groups.items():
        # Section group row (merged)
        sec_row = tbl.add_row()
        # Merge all 6 cells in this row
        sec_row.cells[0].merge(sec_row.cells[5])
        cell = sec_row.cells[0]
        _set_cell_bg(cell, C_SEC_ROW)
        _set_cell_borders(cell, "AAAAAA", "4")
        _set_cell_margins(cell, 60, 60, 100, 100)
        _cell_write(cell, section_name, bold=True, font_size=9,
                    color=C_NAVY, align=WD_ALIGN_PARAGRAPH.LEFT)

        for item in section_items:
            is_mandatory = str(item.get("mandatory", "")).lower() in ("true", "yes", "1")
            status = "MANDATORY" if is_mandatory else item.get("category", "CONDITIONAL")
            risk = _infer_risk(item.get("document_name", ""), is_mandatory, status)

            bg = C_ROW_ODD if row_num % 2 == 1 else C_ROW_EVEN
            row = tbl.add_row()

            # SL
            _set_cell_bg(row.cells[0], bg)
            _set_cell_borders(row.cells[0])
            _set_cell_margins(row.cells[0])
            _cell_write(row.cells[0], str(item.get("doc_id", row_num)),
                        bold=True, font_size=8.5, color=C_NAVY,
                        align=WD_ALIGN_PARAGRAPH.CENTER)

            # Document name
            _set_cell_bg(row.cells[1], bg)
            _set_cell_borders(row.cells[1])
            _set_cell_margins(row.cells[1])
            _cell_write(row.cells[1], item.get("document_name", ""), bold=True, font_size=8.5, color=C_NAVY)

            # Detail requirements
            _set_cell_bg(row.cells[2], bg)
            _set_cell_borders(row.cells[2])
            _set_cell_margins(row.cells[2])
            detail = item.get("description", "") or item.get("notes", "") or "—"
            _cell_write(row.cells[2], detail[:300], font_size=8, color="333333")

            # Reference / Format
            _set_cell_bg(row.cells[3], bg)
            _set_cell_borders(row.cells[3])
            _set_cell_margins(row.cells[3])
            ref_text = item.get("format", "") or ""
            copies = item.get("copies_required", "")
            if copies:
                ref_text = f"{ref_text}\n{copies}" if ref_text else copies
            _cell_write(row.cells[3], ref_text[:150], font_size=8, color="333333")

            # Status
            _set_cell_bg(row.cells[4], _status_fill(status))
            _set_cell_borders(row.cells[4])
            _set_cell_margins(row.cells[4])
            _cell_write(row.cells[4], status, bold=True, font_size=8,
                        color=_status_color(status), align=WD_ALIGN_PARAGRAPH.CENTER)

            # Rejection Risk
            _set_cell_bg(row.cells[5], bg)
            _set_cell_borders(row.cells[5])
            _set_cell_margins(row.cells[5])
            _cell_write(row.cells[5], risk, bold=(risk == "Critical"), font_size=8,
                        color=_risk_color(risk), align=WD_ALIGN_PARAGRAPH.CENTER)

            row_num += 1


def _detect_section(item: dict, env_num: int) -> str:
    """Map an envelope item to a logical section group name."""
    name = (item.get("document_name") or "").lower()
    desc = (item.get("description") or "").lower()
    combined = name + " " + desc

    if any(x in combined for x in ["letter of bid", "submission letter", "authorization", "authorisation", "integrity"]):
        return "SECTION A: SUBMISSION FORMS & AUTHORISATIONS"
    if any(x in combined for x in ["bid security", "tender security", "bank guarantee"]):
        return "SECTION B: BID SECURITY"
    if any(x in combined for x in ["eligibility", "trade license", "tin", "vat", "affidavit", "litigation"]):
        return "SECTION C: ELIGIBILITY & LEGAL DOCUMENTS"
    if any(x in combined for x in ["experience", "completion certif", "end user", "contract award"]):
        return "SECTION D: QUALIFICATION — EXPERIENCE"
    if any(x in combined for x in ["financial", "liquid asset", "turnover", "bank statement", "audited", "annual"]):
        if env_num == 1:
            return "SECTION E: QUALIFICATION — FINANCIAL CAPACITY"
        return "SECTION B: FINANCIAL STATEMENTS"
    if any(x in combined for x in ["personnel", "cv", "curriculum", "key staff", "engineer"]):
        return "SECTION F: KEY PERSONNEL"
    if any(x in combined for x in ["equipment", "tool", "plant"]):
        return "SECTION G: EQUIPMENT & TOOLS"
    if any(x in combined for x in ["oem", "manufacturer", "type test", "iso", "test certif", "catalogue", "brochure"]):
        return "SECTION H: OEM / MANUFACTURER DOCUMENTS"
    if any(x in combined for x in ["technical proposal", "work plan", "methodology", "schedule", "bar chart",
                                    "organogram", "compliance sheet", "specification", "conformity"]):
        return "SECTION I: TECHNICAL PROPOSAL"
    if any(x in combined for x in ["price schedule", "schedule no.", "bill of quantities", "boq", "lcos", "advance payment security"]):
        return "SECTION B: PRICE SCHEDULES"
    if any(x in combined for x in ["jv", "joint venture", "subcontract"]):
        return "SECTION J: CONDITIONAL DOCUMENTS (JV / SUBCONTRACT)"
    if env_num == 2:
        return "SECTION A: FINANCIAL OFFER DOCUMENTS"
    return "SECTION A: SUBMISSION FORMS & AUTHORISATIONS"


# ─── Section 3: Envelope-01 ───────────────────────────────────────────────────

def _add_envelope1_section(doc: Document, env: dict):
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    e1_count = len(env.get("envelope_1", []))
    _section_heading(doc,
                     f"ENVELOPE-01: TECHNICAL OFFER — COMPLETE CONTENTS LIST  ({e1_count} Items)",
                     color=C_BLUE_HDR, font_size=14, space_before=12)

    _para(doc,
          f"The following {e1_count} items must be compiled into the Technical Offer. "
          "Items are arranged by logical grouping. ALL items marked MANDATORY must be included — "
          "missing any Critical or High item will result in non-responsiveness.",
          font_size=9, color="000000", space_before=2, space_after=4)

    _warning(doc, "CRITICAL RULE: The Financial Offer (price schedules, financial statements) "
                  "must NOT be placed inside Envelope-01. Doing so will disqualify the bid at opening.")

    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    _add_envelope_table(doc, env.get("envelope_1", []), env_num=1)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ─── Section 4: Envelope-02 ───────────────────────────────────────────────────

def _add_envelope2_section(doc: Document, env: dict):
    e2_count = len(env.get("envelope_2", []))
    _section_heading(doc,
                     f"ENVELOPE-02: FINANCIAL OFFER — COMPLETE CONTENTS LIST  ({e2_count} Items)",
                     color=C_GREEN_HDR, font_size=14, space_before=14)

    _para(doc,
          f"The following {e2_count} items must be compiled into the Financial Offer. "
          "This envelope is kept sealed until after Technical Offer evaluation is approved. "
          "Only technically responsive bidders' Financial Offers will be opened.",
          font_size=9, color="000000", space_before=2, space_after=4)

    _warning(doc, "CRITICAL RULE: The Bid Security (EMD) belongs in Envelope-01 (Technical Offer) "
                  "NOT in the Financial Offer. The Financial Offer contains ONLY price data and "
                  "financial standing documents as listed below.")

    _warning(doc, "FIXED PRICE: All prices must be FIXED for the duration of the contract. "
                  "Any adjustable price quotation will be treated as non-responsive and rejected.")

    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    _add_envelope_table(doc, env.get("envelope_2", []), env_num=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ─── Section 5: Compliance Checklist ──────────────────────────────────────────

def _add_compliance_section(doc: Document, comp: dict):
    checklist = comp.get("checklist", [])
    if not checklist:
        return

    _section_heading(doc, "COMPLIANCE CHECKLIST — REQUIREMENTS ANALYSIS", color=C_NAVY, space_before=14)

    # Legend
    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(2)
    lp.paragraph_format.space_after = Pt(6)
    for symbol, label, color in [
        ("■ ", "MANDATORY", C_RED),
        ("   ■ ", "CONDITIONAL", C_ORANGE),
        ("   ■ ", "OPTIONAL", "375623"),
    ]:
        r = lp.add_run(symbol)
        r.font.color.rgb = RGBColor.from_string(color)
        r.font.size = Pt(10)
        r2 = lp.add_run(label)
        r2.bold = True
        r2.font.size = Pt(8.5)
        r2.font.name = "Arial"

    # Group by topic
    topics: dict[str, list] = {}
    topic_map = {
        "Eligibility": "ELIGIBILITY REQUIREMENTS",
        "Financial": "FINANCIAL CAPABILITY",
        "Technical": "TECHNICAL CAPABILITY & EXPERIENCE",
        "Legal": "LEGAL & STATUTORY",
        "Bid Security": "BID SECURITY & PERFORMANCE SECURITY",
        "Performance": "BID SECURITY & PERFORMANCE SECURITY",
        "Specification": "TECHNICAL SPECIFICATIONS & STANDARDS",
        "Environmental": "ENVIRONMENTAL & SOCIAL COMPLIANCE",
        "Administrative": "ADMINISTRATIVE & SUBMISSION",
        "Submission": "ADMINISTRATIVE & SUBMISSION",
    }
    for item in checklist:
        sec = item.get("section_reference") or item.get("topic") or "General"
        grp = "GENERAL REQUIREMENTS"
        for keyword, label in topic_map.items():
            if keyword.lower() in sec.lower() or keyword.lower() in (item.get("requirement") or "").lower():
                grp = label
                break
        topics.setdefault(grp, []).append(item)

    tbl = doc.add_table(rows=0, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = tbl.add_row()
    for cell, text in zip(hdr.cells,
                          ["ID", "Requirement", "Document Required", "Ref / Section", "Status"]):
        _set_cell_bg(cell, C_BLUE_HDR)
        _set_cell_borders(cell, "FFFFFF", "4")
        _set_cell_margins(cell)
        _cell_write(cell, text, bold=True, font_size=8.5, color=C_WHITE,
                    align=WD_ALIGN_PARAGRAPH.CENTER)

    row_num = 1
    for topic_name, items in topics.items():
        sec_row = tbl.add_row()
        sec_row.cells[0].merge(sec_row.cells[4])
        cell = sec_row.cells[0]
        _set_cell_bg(cell, C_SEC_ROW)
        _set_cell_borders(cell, "AAAAAA", "4")
        _set_cell_margins(cell, 60, 60, 100, 100)
        _cell_write(cell, topic_name, bold=True, font_size=9, color=C_NAVY)

        for item in items:
            cat = str(item.get("category") or "MANDATORY").upper()
            bg = C_ROW_ODD if row_num % 2 == 1 else C_ROW_EVEN
            row = tbl.add_row()
            for cell in row.cells:
                _set_cell_bg(cell, bg)
                _set_cell_borders(cell)
                _set_cell_margins(cell)

            _cell_write(row.cells[0], item.get("requirement_id", ""),
                        bold=True, font_size=8, color=C_NAVY,
                        align=WD_ALIGN_PARAGRAPH.CENTER)

            req_text = item.get("requirement", "")
            cond = item.get("condition", "")
            if cond:
                req_text += f"  [Condition: {cond}]"
            _cell_write(row.cells[1], req_text, font_size=8, color="222222")
            _cell_write(row.cells[2], item.get("document_required") or "—", font_size=8, color="333333")
            _cell_write(row.cells[3], item.get("section_reference") or "—", font_size=7.5, color=C_GREY_TXT)
            _set_cell_bg(row.cells[4], _status_fill(cat))
            _cell_write(row.cells[4], cat, bold=True, font_size=8,
                        color=_status_color(cat), align=WD_ALIGN_PARAGRAPH.CENTER)
            row_num += 1

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ─── Section 6: OEM Checklist ─────────────────────────────────────────────────

def _add_oem_section(doc: Document, oem: dict):
    requirements = oem.get("oem_requirements", [])
    if not requirements:
        return

    _section_heading(doc, "OEM / MANUFACTURER DOCUMENT CHECKLIST", color=C_NAVY, space_before=14)

    # Approved makes
    makes = oem.get("approved_makes") or []
    if makes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run("Approved Makes / Brands:  ")
        r1.bold = True; r1.font.size = Pt(8.5); r1.font.name = "Arial"
        r1.font.color.rgb = RGBColor.from_string(C_NAVY)
        r2 = p.add_run(", ".join(str(m) for m in makes))
        r2.font.size = Pt(8.5); r2.font.name = "Arial"

    local = oem.get("local_agent_requirements") or ""
    if local:
        _para(doc, f"Local Agent / Dealer Requirements:  {str(local)[:200]}",
              font_size=8.5, color="000000", space_before=1, space_after=4)

    tbl = doc.add_table(rows=0, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = tbl.add_row()
    for cell, text in zip(hdr.cells,
                          ["ID", "Equipment Item", "Requirement Type",
                           "Document Required", "Issuing Party", "Mandatory"]):
        _set_cell_bg(cell, C_BLUE_HDR)
        _set_cell_borders(cell, "FFFFFF", "4")
        _set_cell_margins(cell)
        _cell_write(cell, text, bold=True, font_size=8.5, color=C_WHITE,
                    align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, item in enumerate(requirements):
        is_mandatory = str(item.get("mandatory", "")).lower() in ("true", "yes", "1")
        bg = C_ROW_ODD if i % 2 == 1 else C_ROW_EVEN
        row = tbl.add_row()
        for cell in row.cells:
            _set_cell_bg(cell, bg)
            _set_cell_borders(cell)
            _set_cell_margins(cell)

        _cell_write(row.cells[0], item.get("item_id", str(i + 1)),
                    bold=True, font_size=8, color=C_NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_write(row.cells[1], item.get("equipment_item", ""), bold=True, font_size=8, color=C_NAVY)
        _cell_write(row.cells[2], item.get("oem_requirement_type", ""), font_size=8, color="333333")
        _cell_write(row.cells[3], item.get("required_document") or "—", font_size=8, color="333333")
        _cell_write(row.cells[4], item.get("issuing_party") or "—", font_size=8, color="333333")
        _set_cell_bg(row.cells[5], C_MAND_FILL if is_mandatory else C_OPT_FILL)
        _cell_write(row.cells[5], "YES" if is_mandatory else "No",
                    bold=is_mandatory, font_size=8,
                    color=C_RED if is_mandatory else "375623",
                    align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ─── Section 7: Quick Cross-Reference ────────────────────────────────────────

def _add_crossref_section(doc: Document, env: dict):
    _section_heading(doc, "QUICK CROSS-REFERENCE: WHICH ENVELOPE?",
                     color=C_NAVY, space_before=14)
    _para(doc, "Use this table as a final check before sealing envelopes.",
          font_size=9, color="000000", space_before=2, space_after=5)

    e1_names = {item.get("document_name", "") for item in env.get("envelope_1", [])}
    e2_names = {item.get("document_name", "") for item in env.get("envelope_2", [])}
    all_docs = []
    seen = set()
    for item in env.get("envelope_1", []):
        n = item.get("document_name", "")
        if n not in seen:
            all_docs.append(n)
            seen.add(n)
    for item in env.get("envelope_2", []):
        n = item.get("document_name", "")
        if n not in seen:
            all_docs.append(n)
            seen.add(n)

    tbl = doc.add_table(rows=0, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = tbl.add_row()
    for cell, text in zip(hdr.cells,
                          ["Document / Item", "Envelope-01 (Technical)", "Envelope-02 (Financial)"]):
        _set_cell_bg(cell, C_BLUE_HDR)
        _set_cell_borders(cell, "FFFFFF", "4")
        _set_cell_margins(cell)
        _cell_write(cell, text, bold=True, font_size=8.5, color=C_WHITE,
                    align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, doc_name in enumerate(all_docs):
        in_e1 = doc_name in e1_names
        in_e2 = doc_name in e2_names
        bg = C_ROW_EVEN if i % 2 == 0 else C_ROW_ODD
        row = tbl.add_row()
        _set_cell_bg(row.cells[0], bg)
        _set_cell_borders(row.cells[0])
        _set_cell_margins(row.cells[0])
        _cell_write(row.cells[0], doc_name, font_size=8.5, color="222222")

        for j, (cell, in_env) in enumerate([(row.cells[1], in_e1), (row.cells[2], in_e2)]):
            _set_cell_bg(cell, C_ENV1_HDR if in_env and j == 0 else (C_ENV2_HDR if in_env else bg))
            _set_cell_borders(cell)
            _set_cell_margins(cell)
            _cell_write(cell,
                        "✔ YES" if in_env else "✘ NO",
                        bold=in_env, font_size=8.5,
                        color="1F497D" if in_env and j == 0 else ("375623" if in_env else C_GREY_TXT),
                        align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ─── Footer note ─────────────────────────────────────────────────────────────

def _add_footer_note(doc: Document, meta: dict):
    ref = meta.get("tender_reference") or "this tender"
    entity = meta.get("procuring_entity") or "the Procuring Entity"
    _para(doc,
          f"NOTE: This checklist is prepared based on tender document {ref} as issued by {entity}. "
          "Bidders should verify against any addenda issued subsequent to the tender publication date. "
          "This document does not constitute legal or procurement advice.",
          font_size=8, italic=True, color=C_GREY_NOTE,
          space_before=6, space_after=4)


# ─── Main ─────────────────────────────────────────────────────────────────────

def generate_report(timestamp: str | None = None, output_path: str | None = None) -> Path:
    print("[Report] Loading JSON outputs...")
    meta  = _load(METADATA_OUTPUT_DIR,   timestamp)
    env   = _load(ENVELOPE_OUTPUT_DIR,   timestamp)
    comp  = _load(COMPLIANCE_OUTPUT_DIR, timestamp)
    oem   = _load(OEM_OUTPUT_DIR,        timestamp)

    print(f"  Metadata    : {len(meta)} fields")
    print(f"  Compliance  : {len(comp.get('checklist', []))} requirements")
    print(f"  OEM         : {len(oem.get('oem_requirements', []))} items")
    print(f"  Envelope 1  : {len(env.get('envelope_1', []))} docs")
    print(f"  Envelope 2  : {len(env.get('envelope_2', []))} docs")

    doc = Document()

    # Page setup — A4, tight margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Default style
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)

    print("[Report] Building document...")
    _add_title_block(doc, meta)
    _add_submission_overview(doc, meta, env)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    _add_metadata_section(doc, meta)
    doc.add_page_break()
    _add_envelope1_section(doc, env)
    doc.add_page_break()
    _add_envelope2_section(doc, env)
    doc.add_page_break()
    _add_compliance_section(doc, comp)
    doc.add_page_break()
    _add_oem_section(doc, oem)
    doc.add_page_break()
    _add_crossref_section(doc, env)
    _add_footer_note(doc, meta)

    # Output path
    if output_path:
        out = Path(output_path)
    else:
        ts = timestamp or datetime.now().strftime("%Y%m%d_%H%M%S")
        project = (meta.get("project_name") or "tender")
        # Clean filename
        safe = "".join(c if c.isalnum() or c in "_ -" else "_" for c in project)[:40].strip()
        out = OUTPUTS_DIR / f"{ts}_BidComplianceReport_{safe}.docx"

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"[Report] Saved -> {out}")
    return out


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate CEO-ready Word compliance report")
    parser.add_argument("--timestamp", "-t", help="Run timestamp e.g. 20260611_153155")
    parser.add_argument("--output",    "-o", help="Output .docx path")
    args = parser.parse_args()
    path = generate_report(timestamp=args.timestamp, output_path=args.output)
    print(f"\nDone!  Open: {path}\n")
