"""Tests for document reading tools."""

import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock

from src.tools.document_reader import read_document


def test_unsupported_file_type():
    with pytest.raises(ValueError, match="Unsupported file type"):
        read_document("document.txt")


def test_file_not_found():
    with pytest.raises(FileNotFoundError):
        read_document("nonexistent.pdf")


def test_read_docx(tmp_path):
    """Test DOCX reading with a minimal document."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Test tender document")
    doc.add_paragraph("Project Name: Highway Construction")
    docx_path = tmp_path / "test.docx"
    doc.save(str(docx_path))

    text, file_type = read_document(str(docx_path))
    assert file_type == "docx"
    assert "Test tender document" in text
    assert "Highway Construction" in text
